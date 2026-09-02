# -*- coding: utf-8 -*-
"""
report_builder.py
==================

"Şahin/Güvercin Paneli" uygulamasının verilerini kullanarak, tek bir PPK
kararı için genişletilmiş bir Word (.docx) bilgi notu üretir.

TASARIM İLKESİ — bu modül DB'ye / modele DOKUNMAZ
---------------------------------------------------
Bu dosya hiçbir yerde Supabase, EVDS ya da transformer modeli çağırmaz.
Tüm veriler dışarıdan (önceden hesaplanmış DataFrame'ler olarak) verilir.
Nedeni: `utils.py`'deki ağır fonksiyonlar (`sync_cache`, `calculate_ai_trend_series`,
model indirme vb.) zaten Streamlit uygulaması içinde, önbelleği doldurmak için
çalıştırılıyor. Rapor modülünün işi SADECE o önbellekten / önceden hesaplanmış
tablolardan okunabilir veriyi bir araya getirip biçimlendirmektir. Bu ayrım
sayesinde:
  - Rapor üretimi hızlıdır (model çalışmaz),
  - Modül Streamlit dışında (örn. bir test scriptinde) de çağrılabilir,
  - `app.py` içinde tek satırlık bir entegrasyon yeterlidir (bkz. app.py'deki
    "📄 Rapor" sekmesi — generate_full_report_for_period()).

Kullanım (Streamlit içinden, canlı veriyle):
    from report_builder import generate_full_report_for_period
    path = generate_full_report_for_period(donem="2026-07", analyst_note="...")

Kullanım (bağımsız / test, enjekte edilmiş veriyle):
    from report_builder import build_report
    build_report(df_logs, df_events, df_market, abg_df, ai_df, df_sent,
                  donem="2026-07", model_pack=model_pack, out_path="rapor.docx")

Bağımlılıklar (requirements.txt'e eklenmeli):
    python-docx
    plotly==5.24.1      # DİKKAT: daha yeni plotly (7.x), kaleido>=1 ve dolayısıyla
    kaleido==0.2.1       # ayrı bir Chrome kurulumu ister. Sunucusuz/headless PNG
                          # export için bu ikili en sorunsuz kombinasyondur.
"""

from __future__ import annotations

import io
import os
import tempfile
import datetime as _dt
from typing import Optional

from collections import Counter

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import utils  # noqa: F401  -- proje ile aynı dizinde olmalı


# =============================================================================
# KURUMSAL KİMLİK (TCMB "Rapor Şablonu" stil kılavuzundan) — renkler, yazı tipi
# =============================================================================
# Kaynak: kullanıcının paylaştığı rapor_sablonu.dotx (Renk Paleti (RGB) sayfası +
# "Rapor Tasarımında Uyulacak Kurallar" sayfası). Bu modülün ürettiği rapor artık
# ad-hoc renkler yerine BU paleti kullanır — böylece rapor, TCMB içi diğer
# raporlarla aynı görsel dilde okunur.
BRAND_FONT = "Open Sans"
BRAND_RED = "D50032"        # Kurumsal renk (PANTONE 199C) — tablo başlıkları, vurgular
BRAND_RED_DARK = "9F0024"   # Başlık1 rengi (şablonda accent1 %75 gölge)
BRAND_NAVY = "1C3144"       # Grafik rengi 2
BRAND_GOLD = "F9C213"       # Grafik rengi 3
BRAND_BLUE = "5E8CC6"       # Grafik rengi 4
BRAND_GRAY_BG = "F2F2F2"    # Tablo zebra dolgusu / kutu zemini
BRAND_GRAY_TEXT = "595959"
_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
BRAND_LOGO_PATH = os.path.join(_ASSETS_DIR, "tcmb_logo.png")


def _set_run_font(run, name=BRAND_FONT):
    """Bir run'ın hem Latin hem 'complex script' yazı tipini ayarlar — Word bazen
    Calibri'ye geri düşer çünkü varsayılan tema yazı tipi yalnızca ana Latin
    alanını değiştirir."""
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def _apply_brand_styles(doc):
    """
    Doküman genelinde şablonun yazı tipi/başlık kurallarını uygular:
    Open Sans, Normal 10pt / Başlık1 16pt / Başlık2 13pt / Başlık3 11pt,
    başlıklar siyah-kalın (şablondaki "Rapor Tasarımında Uyulacak Kurallar"
    sayfasındaki punto tablosuyla birebir).
    """
    try:
        normal = doc.styles["Normal"]
        normal.font.name = BRAND_FONT
        normal.font.size = Pt(10)
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), BRAND_FONT)
    except Exception:
        pass

    heading_specs = {  # level: (pt, space_before, space_after)
        1: (16, 12, 6),
        2: (13, 10, 4),
        3: (11, 8, 2),
    }
    for lvl, (pt, before, after) in heading_specs.items():
        try:
            st = doc.styles[f"Heading {lvl}"]
            st.font.name = BRAND_FONT
            st.font.size = Pt(pt)
            st.font.bold = True
            st.font.color.rgb = RGBColor.from_string("1A1A1A")
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
            rPr = st.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), BRAND_FONT)
        except Exception:
            pass


def _set_section_border(paragraph, color=BRAND_RED, sz=10, pos="bottom"):
    """Bir paragrafın altına/üstüne renkli ince bir çizgi ekler (şablondaki
    başlık altı kırmızı çizgi ve kapak çizgisi için)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    edge = OxmlElement(f"w:{pos}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(sz))
    edge.set(qn("w:space"), "4")
    edge.set(qn("w:color"), color)
    pBdr.append(edge)


def _setup_header_footer(doc, title: str, donem: str):
    """
    Şablondaki iç sayfa üstbilgisini ('Rapor Adı | Dönemi [sayfa no]' + ince
    kırmızı çizgi) taklit eder. Kapak sayfasında (ilk sayfa) üstbilgi/altbilgi
    GÖSTERİLMEZ — şablonda da kapakta harici bilgiye yer verilmiyordu.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # ilk sayfa (kapak) üstbilgisi boş kalsın
    _ = section.first_page_header

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = hp.add_run(f"{title}  |  ")
    _set_run_font(r1); r1.font.size = Pt(9); r1.font.color.rgb = RGBColor.from_string("404040")
    r2 = hp.add_run(f"{donem}")
    _set_run_font(r2); r2.font.size = Pt(9); r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(BRAND_RED)
    _set_section_border(hp, color=BRAND_RED, sz=8, pos="bottom")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Şahin/Güvercin Paneli — otomatik üretilmiştir")
    _set_run_font(fr); fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string("8C8C8C")


def _add_cover_page(doc, title: str, donem: str, birim: Optional[str] = None):
    """
    rapor_sablonu.dotx kapak sayfasıyla aynı tipografik hiyerarşiyi kullanır:
    Rapor adı 35pt bütün-büyük-kalın / Hazırlayan birim 16pt bütün-büyük-kalın /
    Dönem 16pt bütün-büyük-kalın KURUMSAL KIRMIZI + TCMB logosu sol altta.
    Şablondaki elmas-desenli fotoğraf kolajı yeniden üretilmedi (karmaşık grafik
    kompozisyonu); onun yerine aynı marka rengiyle sade bir vurgu çizgisi kullanıldı.
    """
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    r = title_p.add_run(title.upper())
    _set_run_font(r); r.bold = True; r.font.size = Pt(35)
    r.font.color.rgb = RGBColor.from_string("1A1A1A")
    _set_section_border(title_p, color=BRAND_RED, sz=18, pos="bottom")
    title_p.paragraph_format.space_after = Pt(18)

    sub_p = doc.add_paragraph()
    r = sub_p.add_run("Genişletilmiş Bilgi Notu")
    _set_run_font(r); r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
    sub_p.paragraph_format.space_after = Pt(28)

    if birim:
        birim_p = doc.add_paragraph()
        r = birim_p.add_run(birim.upper())
        _set_run_font(r); r.bold = True; r.font.size = Pt(16)
        r.font.color.rgb = RGBColor.from_string("1A1A1A")
        birim_p.paragraph_format.space_after = Pt(10)

    donem_p = doc.add_paragraph()
    r = donem_p.add_run(f"DÖNEM: {donem}")
    _set_run_font(r); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BRAND_RED)

    meta_p = doc.add_paragraph()
    r = meta_p.add_run(f"Üretim tarihi: {_dt.date.today().isoformat()}")
    _set_run_font(r); r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(BRAND_GRAY_TEXT)

    for _ in range(10):
        doc.add_paragraph()

    if os.path.exists(BRAND_LOGO_PATH):
        try:
            logo_p = doc.add_paragraph()
            run = logo_p.add_run()
            run.add_picture(BRAND_LOGO_PATH, width=Inches(1.7))
        except Exception:
            pass

    doc.add_page_break()


# =============================================================================
# 0. KÜÇÜK YARDIMCILAR (docx biçimlendirme)
# =============================================================================

def _shade_cell(cell, hex_color: str):
    """Bir tablo hücresini düz renkle doldurur (python-docx bunu public API'de sunmaz)."""
    hex_color = hex_color.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def _tone_hex(diff: float, deadband: Optional[float] = None) -> str:
    """
    utils._tone_rgba ile AYNI eşik ve renk mantığı (kırmızı=şahin #c0392b,
    mavi=güvercin #1f4e9c, nötr=gri), ama docx hücre dolgusu düz renk kabul
    ettiği için alfa karışımı beyaz zemin üzerinden hex'e çevrilmiştir. Uygulamanın
    "🗺️ Ton Haritası" sekmesindeki renklerle raporun renkleri bu sayede birebir
    aynı mantığı izler.
    """
    db = utils.DOC_STANCE_DEADBAND if deadband is None else float(deadband)
    d = float(np.clip(0.0 if pd.isna(diff) else diff, -1.0, 1.0))
    if abs(d) < db:
        return "E9E9EC"  # nötr: düz gri
    t = (abs(d) - db) / max(1e-9, 1.0 - db)
    a = 0.25 + 0.40 * min(1.0, t)
    base = (192, 57, 43) if d > 0 else (31, 78, 156)
    blended = tuple(round(255 * (1 - a) + c * a) for c in base)
    return "".join(f"{c:02X}" for c in blended)


def _shade_run(run, hex_color: str):
    """_shade_cell'in RUN (metin parçası) karşılığı — bir kelime/cümlenin arka
    planını renklendirir. python-docx bunu public API'de sunmaz."""
    hex_color = hex_color.lstrip("#")
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _stance_color_hex(word: str) -> str:
    return {"şahin": "C0392B", "güvercin": "1F4E9C"}.get(word, "7F7F7F")


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("595959")
    return p


def _add_note(doc, text, label="Not:"):
    """Metodolojik uyarı/kısıt kutusu (gri arka planlı tek satırlık tablo)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "F2F2F2")
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    return tbl


def _add_takeaway(doc, text, label="Kısaca:"):
    """
    Her ağır/teknik bölümün EN ÜSTÜNE, o bölümü tek cümlede özetleyen, düz dille
    yazılmış bir "hızlı okuma" kutusu. Amaç: okuyucu teknik detaya (yöntem,
    formül, eşik) girmeden önce "bu bölüm ne anlatıyor, benim için sonucu ne"
    sorusunun cevabını bulsun — detay her zaman ALTINDA, kaybolmadan durmaya
    devam eder. Görsel olarak _add_note'tan (gri/nötr, "dikkat" tonunda) FARKLI
    bir renkte (soluk kurumsal kırmızı) tutulur ki göz otomatik olarak ayırt etsin.
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "FBEAEC")
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    _set_run_font(r1)
    r1.font.color.rgb = RGBColor.from_string(BRAND_RED_DARK)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    _set_run_font(r2)
    r2.font.color.rgb = RGBColor.from_string("1A1A1A")
    return tbl


def _df_to_table(doc, df: pd.DataFrame, header_map: Optional[dict] = None,
                  shade_col: Optional[str] = None, shade_fn=None,
                  col_widths_in=None, max_rows: int = 40, font_size=9):
    """Bir DataFrame'i sade, kenarlıklı bir docx tablosuna çevirir."""
    if df is None or df.empty:
        _add_note(doc, "Bu bölüm için veri bulunamadı.", label="")
        return None
    d = df.head(max_rows).copy()
    cols = list(d.columns)
    labels = [header_map.get(c, c) if header_map else c for c in cols]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Kurumsal tablo kuralı: başlık satırı KURUMSAL KIRMIZI dolgu + beyaz kalın
    # yazı; veri satırları zebra (bir atla bir gri F2F2F2) dolgulu.
    hdr = table.rows[0].cells
    for i, lab in enumerate(labels):
        _set_cell_text(hdr[i], lab, bold=True, size=font_size)
        _shade_cell(hdr[i], BRAND_RED)
        hcell_run = hdr[i].paragraphs[0].runs[0]
        hcell_run.font.color.rgb = RGBColor.from_string("FFFFFF")
        _set_run_font(hcell_run)

    for ridx, (_, row) in enumerate(d.iterrows()):
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            val = row[c]
            if isinstance(val, float):
                txt = f"{val:.3f}" if abs(val) < 10 else f"{val:,.1f}"
            else:
                txt = val
            _set_cell_text(cells[i], txt, size=font_size)
            _set_run_font(cells[i].paragraphs[0].runs[0]) if cells[i].paragraphs[0].runs else None
            if ridx % 2 == 1:
                _shade_cell(cells[i], BRAND_GRAY_BG)
            if shade_col and c == shade_col:
                hexcol = shade_fn(val) if shade_fn else None
                if hexcol:
                    _shade_cell(cells[i], hexcol)

    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    if len(df) > max_rows:
        _add_caption(doc, f"(Tabloda ilk {max_rows} satır gösterilmiştir; toplam {len(df)} satır.)")
    return table


def _tone_gauge_figure(value: Optional[float], regime_label: Optional[str], hyst: float = 25.0):
    """
    Rejimi TEK BAKIŞTA gösteren gösterge (speedometer) grafiği. -100..+100 skalası
    ve renk bantları, utils.postprocess_ai_series_steps'in GERÇEK histerezis eşiğiyle
    (hyst, varsayılan 25.0) birebir tutarlıdır — rastgele/dekoratif bir bant değildir.
    Amaç: sayı yerine "iğnenin nerede durduğunu" gösteren, metin okumadan anlaşılan
    tek bir görsel.
    """
    import plotly.graph_objects as go
    v = 0.0 if value is None or pd.isna(value) else float(np.clip(value, -100, 100))
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=v,
        number={"suffix": "", "font": {"size": 34, "color": "#2c2c2c"}},
        title={"text": f"CB-RoBERTa Rejim Göstergesi<br><span style='font-size:0.65em;color:#666'>"
                        f"{regime_label or '—'}</span>", "font": {"size": 16}},
        gauge={
            "axis": {"range": [-100, 100], "tickwidth": 1, "tickcolor": "#888"},
            "bar": {"color": "#2c3e50", "thickness": 0.28},
            "bgcolor": "white",
            "borderwidth": 0,
            "steps": [
                {"range": [-100, -hyst], "color": "#c9daf0"},   # güvercin bandı
                {"range": [-hyst, hyst], "color": "#eeeeee"},   # nötr bandı
                {"range": [hyst, 100], "color": "#f3cfc8"},     # şahin bandı
            ],
        },
    ))
    fig.update_layout(height=280, margin=dict(t=60, b=10, l=30, r=30))
    return fig


def _add_clean_summary_card(doc, rows: list):
    """
    Uygulamanın Streamlit 'Özet' kartına (kalın etiket + renkli değer + soluk
    '(Önceki: ...)') görsel olarak yakın, KENARLIKSIZ bir liste üretir. Yoğun,
    çok-kolonlu bir tablo yerine tek bakışta okunan bir özet — detaylı yöntem/
    eşik bilgisi ayrı bir tabloda (bkz. hemen altındaki 'Gösterge Detayları').

    rows: [{"label": str, "value": str, "prev": Optional[str], "color": Optional[hex]}]
    """
    if not rows:
        return None
    tbl = doc.add_table(rows=0, cols=2)  # varsayılan stil kenarlıksızdır
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, r in enumerate(rows):
        row = tbl.add_row()
        c0, c1 = row.cells
        if i % 2 == 1:
            _shade_cell(c0, "F5F6F8")
            _shade_cell(c1, "F5F6F8")
        _set_cell_text(c0, r["label"], bold=True, size=11)
        c1.text = ""
        p = c1.paragraphs[0]
        run_v = p.add_run(str(r.get("value", "—")))
        run_v.bold = True
        run_v.font.size = Pt(12)
        if r.get("color"):
            run_v.font.color.rgb = RGBColor.from_string(r["color"])
        if r.get("prev"):
            run_p = p.add_run(f"   (Önceki: {r['prev']})")
            run_p.italic = True
            run_p.font.size = Pt(10)
            run_p.font.color.rgb = RGBColor.from_string("808080")
        c0.width = Inches(2.5)
        c1.width = Inches(3.8)
    return tbl


def _add_sentence_highlight_paragraph(doc, sent_period: pd.DataFrame, deadband: float, show_agent: bool = True):
    """
    utils.sentence_heatmap_html ile AYNI renk mantığını (_tone_hex ↔ _tone_rgba)
    kullanarak, dokümanın TAMAMINI tek akan bir paragraf halinde, her cümlenin
    arka planı kendi tonuna göre boyanmış biçimde yazar. Tablodaki "en şahin /
    en güvercin 6 cümle" özetinin aksine, metnin TAMAMI bağlamıyla birlikte
    okunabilir — uygulamadaki '🗺️ Ton Haritası' sekmesinin docx karşılığı.
    """
    if sent_period is None or sent_period.empty:
        _add_note(doc, "Bu dönem için cümle önbelleği yok, vurgulu metin üretilemedi.")
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for _, r in sent_period.sort_values("sent_idx").iterrows():
        diff = float(r.get("diff", 0.0) or 0.0)
        sent = str(r.get("sentence", "") or "").strip()
        if not sent:
            continue
        run = p.add_run(sent + " ")
        run.font.size = Pt(10)
        _shade_run(run, _tone_hex(diff, deadband))
        agent = r.get("agent_label")
        if show_agent and agent and str(agent).strip():
            run_badge = p.add_run(f" {agent} ")
            run_badge.font.size = Pt(6)
            run_badge.font.color.rgb = RGBColor.from_string("808080")
            run_badge.font.superscript = True

    legend_p = doc.add_paragraph()
    legend_p.paragraph_format.space_before = Pt(2)
    for lbl, val in ((f"şahin (ton ≥ +{deadband:.2f})", 1.0),
                     (f"nötr (|ton| < {deadband:.2f})", 0.0),
                     (f"güvercin (ton ≤ -{deadband:.2f})", -1.0)):
        r1 = legend_p.add_run("  ")
        _shade_run(r1, _tone_hex(val, deadband))
        r2 = legend_p.add_run(f" {lbl}    ")
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor.from_string("595959")
    _add_caption(doc, "Renk yoğunluğu |ton| ile artar; küçük gri üst simge ilgili kesimi gösterir "
                      "(Merkez Bankası / Firmalar / Hanehalkı / Finansal Sektör vb.).")


_TMP_FILES = []


def _fig_to_png(fig, width=1100, height=650, scale=2) -> Optional[str]:
    """plotly Figure -> geçici PNG dosya yolu (kaleido). Başarısızsa None döner
    (rapor üretimini asla düşürmemeli — grafik atlanır, metin devam eder)."""
    if fig is None:
        return None
    try:
        # Isı haritalarındaki (Konu×Ton, Konum×Ton) uzun Türkçe kategori adları,
        # utils.py'nin kendi grafik fonksiyonlarındaki dar sol kenar boşluğuyla
        # (margin l=10) PNG'ye aktarılırken kırpılıyordu — automargin bunu düzeltir.
        fig.update_yaxes(automargin=True)
        fig.update_xaxes(automargin=True)
        fd, path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        fig.write_image(path, width=width, height=height, scale=scale)
        _TMP_FILES.append(path)
        return path
    except Exception as e:
        print(f"[report_builder] Grafik PNG'ye çevrilemedi (atlanıyor): {e}")
        return None


def _mpl_fig_to_png(fig) -> Optional[str]:
    if fig is None:
        return None
    try:
        fd, path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        fig.savefig(path, dpi=170, bbox_inches="tight")
        _TMP_FILES.append(path)
        return path
    except Exception as e:
        print(f"[report_builder] WordCloud PNG'ye çevrilemedi (atlanıyor): {e}")
        return None


def _add_figure(doc, fig, caption=None, width_in=6.3, is_mpl=False, png_width=1100, png_height=650):
    path = _mpl_fig_to_png(fig) if is_mpl else _fig_to_png(fig, width=png_width, height=png_height)
    if not path:
        _add_note(doc, "Bu grafik bu ortamda üretilemedi (kaleido/Chrome eksik olabilir); "
                        "veriler metin/tablo halinde aşağıda mevcut.", label="⚠")
        return
    doc.add_picture(path, width=Inches(width_in))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        _add_caption(doc, caption)


def _cleanup_tmp_files():
    for p in _TMP_FILES:
        try:
            os.remove(p)
        except Exception:
            pass
    _TMP_FILES.clear()


def _fmt_pct(x, digits=1):
    return "—" if x is None or (isinstance(x, float) and pd.isna(x)) else f"%{x:.{digits}f}"


def _fmt_num(x, digits=2, sign=False):
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return "—"
    fmt = f"{{:+.{digits}f}}" if sign else f"{{:.{digits}f}}"
    return fmt.format(x)


# =============================================================================
# 1. VERİYE ÖZGÜ YARDIMCI HESAPLAR
# =============================================================================

def _period_text(df_logs: pd.DataFrame, donem: str) -> str:
    d = df_logs.copy()
    d["Donem"] = pd.to_datetime(d["period_date"]).dt.strftime("%Y-%m")
    row = d[d["Donem"] == str(donem)]
    if row.empty:
        return ""
    return str(row.iloc[0].get("text_content", "") or "")


def _prev_period(df_logs: pd.DataFrame, donem: str) -> Optional[str]:
    d = df_logs.copy()
    d["Donem"] = pd.to_datetime(d["period_date"]).dt.strftime("%Y-%m")
    periods = sorted(d["Donem"].dropna().unique().tolist())
    if donem not in periods:
        return None
    idx = periods.index(donem)
    return periods[idx - 1] if idx > 0 else None


def _hit_rate(pred_df: pd.DataFrame) -> Optional[dict]:
    """Yön isabet oranı (directional hit-rate): tahmin edilen ve gerçekleşen
    delta_bp AYNI YÖNDE mi? utils.py'nin backtest çıktısında bu hesap yok
    (yalnızca MAE/RMSE/R² var) — burada eklendi."""
    if pred_df is None or pred_df.empty:
        return None
    d = pred_df.dropna(subset=["delta_bp", "pred_delta_bp"]).copy()
    if d.empty:
        return None
    d["actual_sign"] = np.sign(d["delta_bp"])
    d["pred_sign"] = np.sign(d["pred_delta_bp"])
    # "Hold" (0 bp) kararlarını ayrı say: yön kıyaslaması yalnızca hamle
    # (hike/cut) olan toplantılar için anlamlıdır.
    moves = d[d["actual_sign"] != 0]
    n = len(moves)
    if n == 0:
        return {"n": 0, "hit_rate": np.nan, "n_total": len(d)}
    hit = (moves["actual_sign"] == moves["pred_sign"]).sum()
    return {"n": int(n), "hit_rate": float(hit) / n, "n_total": int(len(d))}


# =============================================================================
# 2. OTOMATİK METİN SENTEZİ (rapor bunları kendi üretir; LLM'e ihtiyaç yok)
# =============================================================================

def _stance_word(diff, deadband):
    if diff is None or pd.isna(diff):
        return "belirsiz"
    if diff >= deadband:
        return "şahin"
    if diff <= -deadband:
        return "güvercin"
    return "nötr"


def _exec_summary(donem, abg_row, ai_row, ai_prev_row, sent_ozet, flesch, n_sent, flesch_prev, n_sent_prev):
    parts = []
    abg_word = _stance_word((abg_row["abg_index"] - 1.0) if abg_row is not None else None, 0.05)
    ai_word = None
    ai_regime = None
    if ai_row is not None:
        ai_word = _stance_word(ai_row.get("Diff (H-D)"), utils.DOC_STANCE_DEADBAND)
        ai_regime = ai_row.get("AI Rejim") or ai_row.get("Rejim")

    if ai_word and abg_word != "belirsiz":
        agree = "iki yöntem de aynı yönü işaret ediyor" if ai_word == abg_word else \
                "iki yöntem farklı yönlere işaret ediyor (bkz. §3.3 Yöntemler Arası Karşılaştırma)"
        parts.append(
            f"{donem} dönemi PPK metni, sözlük temelli ABG yöntemine göre {abg_word}, "
            f"CB-RoBERTa cümle-bazlı modeline göre {ai_word} olarak sınıflandırılmıştır "
            f"({agree})."
        )
    elif ai_word:
        parts.append(f"{donem} dönemi PPK metni CB-RoBERTa modeline göre {ai_word} olarak sınıflandırılmıştır.")
    elif abg_word != "belirsiz":
        parts.append(f"{donem} dönemi PPK metni ABG yöntemine göre {abg_word} olarak sınıflandırılmıştır.")

    if ai_regime:
        parts.append(f"CB-RoBERTa histerezis rejimi: {ai_regime}.")

    if ai_row is not None and ai_prev_row is not None:
        d_now = ai_row.get("AI Score (EMA)")
        d_prev = ai_prev_row.get("AI Score (EMA)")
        if pd.notna(d_now) and pd.notna(d_prev):
            delta = d_now - d_prev
            yon = "yükseldi (şahinleşti)" if delta > 1 else ("düştü (güvercinleşti)" if delta < -1 else "büyük ölçüde değişmedi")
            parts.append(f"Kalibre edilmiş CB-RoBERTa skoru (EMA) bir önceki döneme göre {yon} "
                         f"({_fmt_num(d_prev,1)} → {_fmt_num(d_now,1)}).")

    if sent_ozet:
        parts.append(
            f"Cümle bazlı dökümde {sent_ozet.get('n','—')} cümlenin "
            f"{sent_ozet.get('n_hawk',0)} tanesi şahin, {sent_ozet.get('n_dove',0)} tanesi güvercin, "
            f"{sent_ozet.get('n_neut',0)} tanesi nötr olarak etiketlenmiştir "
            f"(eşik = ±{utils.DOC_STANCE_DEADBAND:.2f})."
        )

    if flesch is not None:
        d_txt = ""
        if flesch_prev is not None and pd.notna(flesch_prev):
            diff = flesch - flesch_prev
            d_txt = f" (önceki dönem: {flesch_prev:.1f})"
        parts.append(f"Metnin okunabilirlik skoru (Flesch, İngilizce çeviri üzerinden) {flesch:.1f}{d_txt}, "
                     f"cümle sayısı {n_sent}{f' (önceki: {n_sent_prev})' if n_sent_prev else ''} olarak ölçülmüştür.")

    return " ".join(parts) if parts else "Bu dönem için özetleyici veri bulunamadı."


def _reconciliation_text(abg_row, sent_ozet):
    if abg_row is None or not sent_ozet:
        return ("Bu dönem için iki yöntemi karşılaştırmaya yetecek veri bulunamadı.")

    n_match = int(abg_row.get("n_match", 0))
    hawk_abg, dove_abg = int(abg_row.get("hawk_count", 0)), int(abg_row.get("dove_count", 0))
    n_hawk_rb, n_dove_rb, n_neut_rb = sent_ozet.get("n_hawk", 0), sent_ozet.get("n_dove", 0), sent_ozet.get("n_neut", 0)
    n_total_rb = sent_ozet.get("n", 0)

    txt = (
        f"ABG (sözlük temelli) yöntem bu dönemde toplam {n_match} eşleşme buldu "
        f"({hawk_abg} şahin, {dove_abg} güvercin modifikatör); CB-RoBERTa cümle-bazlı model ise "
        f"{n_total_rb} cümlenin tamamını sınıflandırıp {n_hawk_rb} şahin, {n_dove_rb} güvercin, "
        f"{n_neut_rb} nötr cümle buldu."
    )

    if n_match < 5:
        txt += (
            f" ABG yönteminin bu dönemdeki örneklemi ({n_match} eşleşme) düşüktür: sözlük yalnızca "
            f"enflasyon, iktisadi faaliyet ve istihdam etrafındaki sınırlı bir kelime kümesini "
            f"tarar, PPK metninin geri kalanını görmez. Bu nedenle ABG endeksi bu dönem için "
            f"TEK BAŞINA yorumlanmamalı, CB-RoBERTa'nın cümle-bazlı sonuçlarıyla birlikte "
            f"okunmalıdır — düşük eşleşme sayısı endeksin uçlara yapışmasına yol açabilir."
        )
    else:
        txt += " Her iki yöntem de metnin tamamını değil, kendi kapsamındaki ifadeleri ölçtüğü için sayılar birebir örtüşmesi beklenmez; asıl karşılaştırılması gereken YÖN'dür."

    dir_abg = 1 if hawk_abg > dove_abg else (-1 if dove_abg > hawk_abg else 0)
    dir_rb = 1 if n_hawk_rb > n_dove_rb else (-1 if n_dove_rb > n_hawk_rb else 0)
    if dir_abg != 0 and dir_rb != 0:
        if dir_abg == dir_rb:
            txt += " Yön açısından iki yöntem birbiriyle TUTARLIDIR."
        else:
            txt += (" Yön açısından iki yöntem BİRBİRİNDEN AYRIŞIYOR — bu genellikle metnin "
                    "ABG sözlüğünün kapsamadığı temalarda (ör. iletişim/çerçeve dili) yoğunlaştığına "
                    "işaret eder ve tek bir sayıya indirgemek yerine cümle dökümüne (§4) bakılmalıdır.")
    return txt


_FALLBACK_STOPWORDS = {
    "the", "a", "an", "of", "to", "in", "and", "or", "is", "are", "was", "were", "be", "been",
    "being", "on", "at", "by", "for", "with", "as", "that", "this", "it", "its", "which", "will",
    "would", "should", "could", "has", "have", "had", "not", "no", "also", "than", "then", "there",
    "their", "from", "into", "over", "under", "further", "more", "most", "such", "other", "these",
    "those", "based", "same", "may", "can", "if", "so", "but", "while", "when", "where", "s",
}


def _lexical_shift_rows(text_now: str, text_prev: str, top_k: int = 12) -> list:
    """
    Önceki döneme göre ÇOĞALAN/AZALAN/YENİ ORTAYA ÇIKAN/KAYBOLAN kelime ve iki-kelimelik
    ifadeleri bulur. Ham sayım yerine 1000 kelime başına oran (‰) kullanılır — metin
    uzunluğu dönemden döneme değiştiği için ham sayım yanıltıcı olabilir (ör. metin
    kısalırken aynı sayıda geçen bir kelime aslında GÖRECELİ olarak ağırlaşmış olur).
    """
    try:
        from wordcloud import STOPWORDS as _sw
        stop = set(_sw)
    except Exception:
        stop = _FALLBACK_STOPWORDS

    def _freqs(text, ngram):
        if not text:
            return {}
        toks = [t for t in utils.tokenize(utils.normalize_text(text)) if len(t) > 2 and t not in stop]
        items = toks if ngram == 1 else [f"{a} {b}" for a, b in zip(toks, toks[1:])
                                          if a not in stop and b not in stop]
        n = len(items) or 1
        c = Counter(items)
        return {k: v / n * 1000.0 for k, v in c.items()}

    rows = []
    for ngram, label in ((1, "tek kelime"), (2, "iki kelimelik ifade")):
        f_now = _freqs(text_now, ngram)
        f_prev = _freqs(text_prev, ngram)
        for k in set(f_now) | set(f_prev):
            now_v, prev_v = f_now.get(k, 0.0), f_prev.get(k, 0.0)
            if prev_v == 0 and now_v >= 1.0:
                durum = "🆕 Yeni ortaya çıktı"
            elif now_v == 0 and prev_v >= 1.0:
                durum = "❌ Kayboldu"
            elif prev_v > 0 and now_v >= prev_v * 1.6 and (now_v - prev_v) >= 0.5:
                durum = "↑ Güçlendi"
            elif prev_v > 0 and now_v <= prev_v / 1.6 and (prev_v - now_v) >= 0.5:
                durum = "↓ Zayıfladı"
            else:
                continue
            rows.append({"Tür": label, "İfade": k, "Bu dönem (‰)": round(now_v, 2),
                         "Önceki dönem (‰)": round(prev_v, 2), "Değişim": durum})
    rows.sort(key=lambda r: abs(r["Bu dönem (‰)"] - r["Önceki dönem (‰)"]), reverse=True)
    return rows[:top_k]


_DOVE_NOTE_KEYWORDS = ["güvercin", "gevşe", "dovish", "faiz indir", "indirim sinyal", "destekleyici ton"]
_HAWK_NOTE_KEYWORDS = ["şahin", "sıkılaştır", "hawkish", "faiz artış", "sıkı duruş", "sıkı para"]


def _note_consistency_caveat(note: Optional[str], abg_word: str, ai_word: str) -> Optional[str]:
    """
    Yapıştırılan analist notunun yönü, bu rapor için otomatik hesaplanan yönle (ABG/CB-RoBERTa)
    çelişiyorsa uyarı üretir. Amaç: örn. önceki bir toplantıdan kalmış, güncellenmemiş bir notun
    fark edilmeden rapora girmesini engellemek (bkz. §8 kullanım notu). Bu bir dilbilimsel/anlamsal
    analiz DEĞİL, kaba bir anahtar-kelime taramasıdır — kesin bir hüküm olarak okunmamalıdır.
    """
    if not note or not note.strip():
        return None
    low = note.lower()
    has_dove = any(k in low for k in _DOVE_NOTE_KEYWORDS)
    has_hawk = any(k in low for k in _HAWK_NOTE_KEYWORDS)
    if has_dove and has_hawk:
        return None  # her iki yön de geçiyor — muhtemelen nüanslı bir tartışma, bayrak kaldırma
    auto = {w for w in (abg_word, ai_word) if w in ("şahin", "güvercin")}
    if has_dove and auto == {"şahin"}:
        return ("Bu not GÜVERCİN yönlü ifadeler içeriyor, ancak bu dönem için ABG ve/veya CB-RoBERTa "
                "yöntemleri ŞAHİN sonucu veriyor. Not, farklı/eski bir döneme ait olabilir — "
                "hangi metne ve tarihe ait olduğunu kontrol edin.")
    if has_hawk and auto == {"güvercin"}:
        return ("Bu not ŞAHİN yönlü ifadeler içeriyor, ancak bu dönem için ABG ve/veya CB-RoBERTa "
                "yöntemleri GÜVERCİN sonucu veriyor. Not, farklı/eski bir döneme ait olabilir — "
                "hangi metne ve tarihe ait olduğunu kontrol edin.")
    return None


def _position_short(pos_matrix: pd.DataFrame, donem: str) -> Optional[str]:
    """§9 Sonuç için kısa, tek cümlelik konum özeti. _position_narrative'in tam
    metnini "." üzerinden kırpmak (ör. '(+0.41)' içindeki ondalık noktasında)
    cümleyi yanlış yerden kesiyordu; bu yüzden ayrı, sayı İÇERMEYEN bir cümle
    üretiyoruz."""
    if pos_matrix is None or pos_matrix.empty or donem not in pos_matrix.index:
        return None
    row = pos_matrix.loc[donem].dropna()
    if row.empty:
        return None
    row_sorted = row.sort_values(ascending=False)
    return (f"Metin {row_sorted.index[0]} bölümünde en şahin, {row_sorted.index[-1]} "
            f"bölümünde göreli olarak en güvercin/en az şahin tonu taşımaktadır.")


def _position_narrative(pos_matrix: pd.DataFrame, donem: str) -> str:
    if pos_matrix is None or pos_matrix.empty or donem not in pos_matrix.index:
        return "Bu dönem için metin-içi konum verisi bulunamadı."
    row = pos_matrix.loc[donem].dropna()
    if row.empty:
        return "Bu dönemde konum bazlı ton hesaplanamadı (yetersiz cümle sayısı)."
    row_sorted = row.sort_values(ascending=False)
    en_sahin = row_sorted.index[0]
    en_guvercin = row_sorted.index[-1]
    spread = row_sorted.iloc[0] - row_sorted.iloc[-1]
    txt = (
        f"Metin {en_sahin} bölümünde en şahin tonu taşırken ({row_sorted.iloc[0]:+.2f}), "
        f"{en_guvercin} bölümünde göreli olarak en güvercin ya da en az şahin tonu taşımaktadır "
        f"({row_sorted.iloc[-1]:+.2f})."
    )
    if spread >= 0.3:
        txt += (" Bölümler arası fark belirgindir; bu, metnin homojen bir ton yerine bilinçli bir "
                "anlatı yapısı (ör. temkinli/tanımlayıcı bir açılış, sıkı bir çerçeve bloğu, "
                "kararlı bir kapanış) izlediğine işaret edebilir.")
    else:
        txt += " Bölümler arası fark sınırlıdır; ton metin boyunca görece homojendir."
    return txt


def _backtest_narrative(metrics: dict, hit: Optional[dict]) -> str:
    if not metrics or pd.isna(metrics.get("mae", np.nan)):
        return "Backtest için yeterli etiketli gözlem (delta_bp dolu karar) bulunamadı."
    txt = (
        f"TF-IDF (kelime + karakter) + gecikmeli makro/faiz özellikleri ile eğitilen Ridge "
        f"regresyon modeli, yürüyen-pencere (walk-forward) doğrulamada {metrics.get('n','—')} "
        f"karar üzerinde ortalama {metrics.get('mae', float('nan')):.1f} baz puan mutlak hata "
        f"(MAE), {metrics.get('rmse', float('nan')):.1f} baz puan RMSE ve "
        f"R²={metrics.get('r2', float('nan')):.2f} vermiştir."
    )
    if hit and hit.get("n", 0) > 0:
        txt += (
            f" Yalnızca faiz hareketi içeren ({hit['n']} adet hike/cut) kararlara bakıldığında, "
            f"modelin tahmin ettiği YÖN (artış/indirim) gerçekleşenle {hit['hit_rate']*100:.0f}% "
            f"oranında örtüşmüştür — bu, mevcut kod tabanında hesaplanmayan, bu rapor için "
            f"eklenmiş bir isabet-oranı metriğidir."
        )
    else:
        txt += " Bu dönem aralığında yön isabet oranı hesaplamaya yetecek sayıda faiz hareketi (hike/cut) yok."
    txt += (" R² ve MAE, mutlak sapmayı ölçer; hit-rate ise pratikte daha çok sorulan soruyu "
            "(\"metin doğru YÖNÜ mü işaret ediyor\") cevaplar — ikisi birlikte okunmalıdır. "
            "Örneklem küçük olduğu için (walk-forward'da her adım bir öncekilerle eğitilir) "
            "metrikler geniş güven aralıklarıyla yorumlanmalıdır.")
    return txt


def _next_meeting_prediction(model_pack: Optional[dict], donem: str, text_now: str,
                              df_logs: pd.DataFrame) -> Optional[dict]:
    """
    Modelin SIRADAKİ (henüz gerçekleşmemiş ya da bu rapordaki dönemden sonraki)
    PPK kararı için ima ettiği delta_bp'yi hesaplar.

    ÖNEMLİ METODOLOJİK NOKTA: PPK karar metni, kararla AYNI ANDA yayımlanır —
    yani "bir sonraki, henüz yazılmamış" karar metni diye bir şey YOKTUR ve
    hiçbir metin-temelli model bunu gerçek anlamda göremez. Bu fonksiyon bunun
    yerine utils.predict_textasdata_hybrid_cpi'yi, elde EN GÜNCEL YAYIMLANMIŞ
    metni (bu raporun kendi döneminin metnini) "ileriye dönük sözlü yönlendirme"
    (forward guidance) sinyali olarak kullanacak şekilde çağırır — utils.py'deki
    bu fonksiyon zaten tam olarak bunun için tasarlanmış (son bilinen durumun
    gecikmeli/lag özelliklerini alıp verilen metinle birleştiriyor). Sonuç bu
    yüzden bir KEHANET değil, "mevcut iletişim biçimi sürerse model ne bekler"
    sorusunun cevabıdır — aşağıdaki kutuda bu sınırlama açıkça belirtilir.
    """
    if not model_pack or "model" not in model_pack:
        return None
    df_hist = model_pack.get("df_hist")
    if df_hist is None or not isinstance(df_hist, pd.DataFrame) or df_hist.empty:
        return None
    if not text_now or len(text_now.strip()) < 30:
        return None

    pred = utils.predict_textasdata_hybrid_cpi(model_pack, df_hist, text_now)
    if not pred or "pred_delta_bp" not in pred:
        return None
    pred_bp = float(pred["pred_delta_bp"])

    # ampirik %10-%90 aralığı: geçmiş walk-forward hatalarının (gerçekleşen - tahmin)
    # dağılımından — modelin KENDİ geçmiş hata payı, ayrı bir varsayım gerektirmez.
    lo, hi = -50.0, 50.0
    pred_df = model_pack.get("pred_df")
    if pred_df is not None and not pred_df.empty and {"delta_bp", "pred_delta_bp"}.issubset(pred_df.columns):
        resid = (pred_df["delta_bp"] - pred_df["pred_delta_bp"]).dropna()
        if len(resid) >= 5:
            lo, hi = float(np.nanquantile(resid, 0.10)), float(np.nanquantile(resid, 0.90))
    interval = tuple(sorted((pred_bp + lo, pred_bp + hi)))

    threshold = 25.0
    if pred_bp >= threshold:
        yon = "ARTIRIM"
    elif pred_bp <= -threshold:
        yon = "İNDİRİM"
    else:
        yon = "SABİT"

    # bu dönemden SONRAKİ bir karar zaten kayıtlıysa (rapor geçmişe dönük üretildiyse),
    # gerçekleşenle karşılaştırma imkânı var — canlı/ileriye dönük bir tahmin değilse bunu belirt.
    actual_next = None
    actual_next_donem = None
    try:
        d = df_logs.copy()
        d["Donem"] = pd.to_datetime(d["period_date"]).dt.strftime("%Y-%m")
        periods = sorted(d["Donem"].dropna().unique().tolist())
        if donem in periods:
            idx = periods.index(donem)
            if idx < len(periods) - 1:
                actual_next_donem = periods[idx + 1]
                row = d[d["Donem"] == actual_next_donem]
                if not row.empty and pd.notna(row.iloc[0].get("delta_bp")):
                    actual_next = float(row.iloc[0]["delta_bp"])
    except Exception:
        pass

    return {
        "pred_bp": pred_bp, "interval": interval, "yon": yon,
        "actual_next": actual_next, "actual_next_donem": actual_next_donem,
        "is_live_forecast": actual_next_donem is None,
    }


def _conclusion(donem, abg_row, ai_row, sent_ozet, backtest_metrics, hit, pos_narrative_short):
    bits = []
    abg_word = _stance_word((abg_row["abg_index"] - 1.0) if abg_row is not None else None, 0.05)
    ai_word = _stance_word(ai_row.get("Diff (H-D)"), utils.DOC_STANCE_DEADBAND) if ai_row is not None else "belirsiz"
    if ai_word == abg_word and ai_word != "belirsiz":
        bits.append(f"{donem} kararı iki bağımsız yöntemde de aynı yönde ({ai_word}) sınıflandığı için "
                    f"genel duruş okuması güvenilir kabul edilebilir.")
    else:
        bits.append(f"{donem} kararında yöntemler arasında {('tam' if ai_word==abg_word else 'kısmi')} "
                    f"uyum var; tek bir skora değil, cümle bazlı döküme (§4) bakılarak karar verilmelidir.")
    if sent_ozet and sent_ozet.get("n", 0) > 0:
        n_hawk, n_dove, n = sent_ozet["n_hawk"], sent_ozet["n_dove"], sent_ozet["n"]
        if n_hawk > 2 * max(n_dove, 1):
            bits.append("Şahin cümleler güvercin cümlelerin belirgin biçimde üzerindedir; "
                        "iletişimde net bir sıkı-duruş ağırlığı vardır.")
        elif n_dove > 2 * max(n_hawk, 1):
            bits.append("Güvercin cümleler şahin cümlelerin belirgin biçimde üzerindedir; "
                        "iletişimde gevşeme sinyali ağır basmaktadır.")
        else:
            bits.append("Şahin ve güvercin cümle sayıları birbirine yakındır; metin karışık/dengeli bir "
                        "ton taşımaktadır — tek yönlü bir okuma yapmak yanıltıcı olabilir.")
    if backtest_metrics and not pd.isna(backtest_metrics.get("r2", np.nan)):
        r2 = backtest_metrics["r2"]
        guven = "düşük" if r2 < 0.1 else ("orta" if r2 < 0.4 else "yüksek")
        bits.append(f"Metin-temelli faiz tahmin modelinin tarihsel açıklama gücü ({guven}, R²={r2:.2f}) "
                    f"göz önünde bulundurulduğunda, bu raporun sinyalleri YÖNLENDİRİCİ değil "
                    f"DESTEKLEYİCİ bilgi olarak kullanılmalıdır.")
    if pos_narrative_short:
        bits.append(pos_narrative_short)
    return " ".join(bits)


# =============================================================================
# 3. ANA FONKSİYON
# =============================================================================

def build_report(
    df_logs: pd.DataFrame,
    df_events: Optional[pd.DataFrame],
    df_market: Optional[pd.DataFrame],
    abg_df: pd.DataFrame,
    ai_df: Optional[pd.DataFrame],
    df_sent: Optional[pd.DataFrame],
    donem: str,
    model_pack: Optional[dict] = None,
    analyst_note: Optional[str] = None,
    out_path: str = "ppk_rapor.docx",
    title: str = "PPK Metni İletişim Analizi",
    birim: Optional[str] = None,
) -> str:
    """
    Tek bir PPK dönemi için genişletilmiş Word raporu üretir.

    Parametrelerin hepsi ÖNCEDEN HESAPLANMIŞ olmalı (bu fonksiyon hiçbir model/DB
    çağrısı yapmaz — bkz. modül başındaki tasarım notu). `donem` "YYYY-MM" biçiminde
    olmalı ve df_logs/abg_df/ai_df/df_sent içinde bu döneme ait en az bir kayıt
    bulunmalıdır (aksi halde ilgili bölümler "veri yok" notuyla atlanır, rapor
    üretimi HATA VERMEZ).
    """
    _TMP_FILES.clear()
    doc = Document()
    _apply_brand_styles(doc)

    # Sayfa kenar boşlukları biraz daraltılsın — tablolar/grafikler için yer açar
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---- Hazırlık: dönem verileri ------------------------------------------
    df_logs = df_logs.copy()
    df_logs["period_date"] = pd.to_datetime(df_logs["period_date"], errors="coerce")
    df_logs["Donem"] = df_logs["period_date"].dt.strftime("%Y-%m")

    text_now = _period_text(df_logs, donem)
    prev_donem = _prev_period(df_logs, donem)
    text_prev = _period_text(df_logs, prev_donem) if prev_donem else ""

    abg_df = abg_df if abg_df is not None else pd.DataFrame()
    abg_row = None
    abg_row_prev = None
    if not abg_df.empty:
        m = abg_df[abg_df["Donem"] == donem]
        abg_row = m.iloc[0] if not m.empty else None
        if prev_donem:
            mp = abg_df[abg_df["Donem"] == prev_donem]
            abg_row_prev = mp.iloc[0] if not mp.empty else None

    ai_row = None
    ai_row_prev = None
    if ai_df is not None and not ai_df.empty:
        m = ai_df[ai_df["Dönem"] == donem]
        ai_row = m.iloc[0] if not m.empty else None
        if prev_donem:
            mp = ai_df[ai_df["Dönem"] == prev_donem]
            ai_row_prev = mp.iloc[0] if not mp.empty else None

    df_sent = df_sent if df_sent is not None else pd.DataFrame()
    sent_period = df_sent[df_sent["Donem"] == donem].copy() if not df_sent.empty else pd.DataFrame()
    sahin_df, guvercin_df, sent_ozet = (pd.DataFrame(), pd.DataFrame(), {})
    if not df_sent.empty:
        sahin_df, guvercin_df, sent_ozet = utils.top_sentences(donem, k=6, df_sent=df_sent)

    flesch_now = utils.calculate_flesch_reading_ease(text_now) if text_now else None
    flesch_prev = utils.calculate_flesch_reading_ease(text_prev) if text_prev else None
    n_sent_now = len(sent_period) if not sent_period.empty else (
        len(utils.split_sentences_nlp(utils.normalize_text(text_now))) if text_now else None)
    n_sent_prev = None
    if prev_donem and not df_sent.empty:
        sp = df_sent[df_sent["Donem"] == prev_donem]
        n_sent_prev = len(sp) if not sp.empty else None

    # =========================================================================
    # KAPAK + İÇ SAYFA ÜSTBİLGİSİ (TCMB rapor şablonuyla uyumlu)
    # =========================================================================
    _add_cover_page(doc, title, donem, birim=birim)
    _setup_header_footer(doc, title, donem)

    # =========================================================================
    # 1. YÖNETİCİ ÖZETİ
    # =========================================================================
    _add_heading(doc, "1. Yönetici Özeti", level=1)
    doc.add_paragraph(_exec_summary(donem, abg_row, ai_row, ai_row_prev, sent_ozet,
                                     flesch_now, n_sent_now, flesch_prev, n_sent_prev))

    # ---- temiz özet kartı (uygulamadaki 'Özet' görünümüne yakın) ----
    _abg_word = _stance_word((abg_row["abg_index"] - 1.0) if abg_row is not None else None, 0.05)
    _abg_word_prev = _stance_word((abg_row_prev["abg_index"] - 1.0) if abg_row_prev is not None else None, 0.05)
    _ai_word = _stance_word(ai_row.get("Diff (H-D)"), utils.DOC_STANCE_DEADBAND) if ai_row is not None else "belirsiz"
    _ai_word_prev = _stance_word(ai_row_prev.get("Diff (H-D)"), utils.DOC_STANCE_DEADBAND) if ai_row_prev is not None else "belirsiz"
    _ai_regime = str(ai_row.get("AI Rejim", ai_row.get("Rejim", "—"))) if ai_row is not None else None
    _ai_regime_prev = str(ai_row_prev.get("AI Rejim", ai_row_prev.get("Rejim", "—"))) if ai_row_prev is not None else None

    card_rows = []
    if abg_row is not None:
        card_rows.append({"label": "Apel, Blix-Grimaldi (sözlük temelli endeks)",
                          "value": _abg_word.capitalize(), "prev": _abg_word_prev.capitalize() if abg_row_prev is not None else None,
                          "color": _stance_color_hex(_abg_word)})
    if ai_row is not None:
        card_rows.append({"label": "CB-RoBERTa (cümle-bazlı model — ton)",
                          "value": _ai_word.capitalize(), "prev": _ai_word_prev.capitalize() if ai_row_prev is not None else None,
                          "color": _stance_color_hex(_ai_word)})
        card_rows.append({"label": "CB-RoBERTa (rejim — histerezis/EMA)",
                          "value": _ai_regime or "—", "prev": _ai_regime_prev,
                          "color": None})
    if flesch_now is not None:
        card_rows.append({"label": "Okunabilirlik skoru (Flesch)",
                          "value": f"{flesch_now:.2f}", "prev": f"{flesch_prev:.2f}" if flesch_prev is not None else None,
                          "color": None})
    if n_sent_now is not None:
        card_rows.append({"label": "Cümle sayısı",
                          "value": str(n_sent_now), "prev": str(n_sent_prev) if n_sent_prev else None,
                          "color": None})
    _add_clean_summary_card(doc, card_rows)

    # ---- gösterge (speedometer) — metne bakmadan "iğnenin nerede durduğunu" gösterir ----
    if ai_row is not None:
        try:
            _gauge_fig = _tone_gauge_figure(ai_row.get("AI Score (EMA)"), _ai_regime, hyst=25.0)
            _add_figure(doc, _gauge_fig, width_in=3.4, png_width=520, png_height=380,
                       caption="İğne, kalibre edilmiş EMA skorunu gösterir; renkli bantlar rejim eşiğini "
                               "(±25) işaretler — bkz. §2.2.")
        except Exception as e:
            _add_note(doc, f"Gösterge grafiği üretilemedi: {e}")

    # ---- özet tablo (detaylı: yöntem/ölçek/eşik) ----
    # Her göstergenin NE ölçtüğü ve hangi ölçek/eşik üzerinden okunduğu, değerin
    # kendisiyle birlikte gösterilir — aksi halde (ör. "ABG=1.00" ile "CB-RoBERTa
    # ton=+0.26" yan yana) okuyucu iki sayının aynı şeyi mi ölçtüğünü sanabilir.
    # (Yukarıdaki temiz kart hızlı okuma için; bu tablo merak edenler için detaydır.)
    _add_heading(doc, "Gösterge Detayları (ölçek, eşik, yöntem notu)", level=2)
    summary_rows = []
    if abg_row is not None:
        summary_rows.append({
            "Gösterge": "ABG (2019) — yumuşatılmış endeks",
            "Ölçtüğü şey": "Sözlük/kural tabanlı şahin-güvercin kelime dengesi (tüm metne bakar)",
            "Ölçek / Eşik": "0–2  (1.00 = nötr; >1 şahin, <1 güvercin)",
            "Değer": f"{abg_row['abg_index']:.3f}  (ham: {abg_row.get('abg_index_raw', float('nan')):.3f})",
            "Önceki dönem": f"{abg_row_prev['abg_index']:.3f}" if abg_row_prev is not None else "—",
            "Not": f"n_match={int(abg_row.get('n_match', 0))} — {'DÜŞÜK ÖRNEKLEM, ihtiyatlı yorumlayın' if abg_row.get('n_match', 0) < 5 else 'yeterli örneklem'}",
        })
    if ai_row is not None:
        summary_rows.append({
            "Gösterge": "CB-RoBERTa — ton (Diff H-D)",
            "Ölçtüğü şey": "Cümle-bazlı modelin bu dönemin HAM tonu (P(Şahin)−P(Güvercin), karar cümlesi ağırlıklı)",
            "Ölçek / Eşik": f"-1..+1  (±{utils.DOC_STANCE_DEADBAND:.2f} = şahin/nötr/güvercin eşiği)",
            "Değer": f"{ai_row.get('Diff (H-D)', float('nan')):+.3f}  ({ai_row.get('Duruş','—')})",
            "Önceki dönem": f"{ai_row_prev.get('Diff (H-D)', float('nan')):+.3f}" if ai_row_prev is not None else "—",
            "Not": "cümle-bazlı, karar-ağırlıklı kanonik sinyal",
        })
        summary_rows.append({
            "Gösterge": "CB-RoBERTa — rejim (histerezis, EMA)",
            "Ölçtüğü şey": "Dönemler arası gürültüyü süzen, yumuşatılmış GENEL İLETİŞİM REJİMİ (tek dönemin tonu değil)",
            "Ölçek / Eşik": "🦅 Şahin / ⚪ Nötr / 🕊️ Güvercin  (histerezis bandı — bkz. §2.2)",
            "Değer": str(ai_row.get("AI Rejim", ai_row.get("Rejim", "—"))),
            "Önceki dönem": str(ai_row_prev.get("AI Rejim", ai_row_prev.get("Rejim", "—"))) if ai_row_prev is not None else "—",
            "Not": "ani rejim sıçramalarını önlemek için yumuşatılmıştır (bkz. §2 Yöntem)",
        })
    if flesch_now is not None:
        summary_rows.append({
            "Gösterge": "Okunabilirlik (Flesch)",
            "Ölçtüğü şey": "Metnin (İngilizce çeviri) ne kadar kolay okunduğu — ŞAHİN/GÜVERCİN yönüyle İLGİSİZDİR",
            "Ölçek / Eşik": "0–100  (yüksek = daha kolay okunur; eşik yok)",
            "Değer": f"{flesch_now:.2f}",
            "Önceki dönem": f"{flesch_prev:.2f}" if flesch_prev is not None else "—",
            "Not": "İngilizce çeviri üzerinden hesaplanır",
        })
    if n_sent_now is not None:
        summary_rows.append({
            "Gösterge": "Cümle sayısı",
            "Ölçtüğü şey": "Metindeki toplam cümle adedi — bu rapordaki TÜM cümle-bazlı sayımların (§4, §5, §6) paydası",
            "Ölçek / Eşik": "sayaç (eşik yok)",
            "Değer": str(n_sent_now),
            "Önceki dönem": str(n_sent_prev) if n_sent_prev else "—",
            "Not": "",
        })
    _df_to_table(doc, pd.DataFrame(summary_rows), font_size=8)
    _add_note(doc, "Bu raporun tüm sayısal göstergeleri, aşağıdaki §2 Yöntem bölümünde tanımlanan "
                   "eşikler ve dönüşümlerle üretilmiştir; farklı bir eşikle yeniden hesaplanırsa "
                   "farklı bir 'şahin/güvercin' etiketi çıkabilir. Yukarıdaki göstergeler FARKLI "
                   "ŞEYLER ölçer (bkz. 'Ölçtüğü şey' kolonu) — sayıca birbirine yakın/uzak olmaları "
                   "tek başına bir tutarlılık/tutarsızlık kanıtı değildir; yöntemler arası "
                   "karşılaştırma için bkz. §3.3.")

    # =========================================================================
    # 2. YÖNTEM
    # =========================================================================
    _add_heading(doc, "2. Yöntem", level=1)
    _add_takeaway(doc,
        "Metnin şahin mi güvercin mi olduğunu İKİ FARKLI YÖNTEMLE ölçüyoruz: basit bir "
        "kelime sayımı (ABG) ve cümle cümle okuyan bir yapay-zekâ modeli (CB-RoBERTa). "
        "İkisi de aynı yöne işaret ederse güven yüksektir; ayrışırlarsa cümle bazlı döküme "
        "(§4) bakılmalıdır. Aşağıdaki alt başlıklar bu yöntemlerin teknik detaylarıdır — "
        "sadece sonucu merak ediyorsanız §1'e dönebilirsiniz.")

    _add_heading(doc, "2.1 ABG (Apel & Blix-Grimaldi, 2019) — sözlük temelli endeks", level=2)
    doc.add_paragraph(
        "Yöntem, metindeki üç konu (enflasyon, iktisadi faaliyet, istihdam) etrafında geçen "
        "belirli çapa kelimelerin (ör. \"inflation\", \"economic activity\", \"employment\") "
        f"±10 kelimelik penceresinde şahin/güvercin yönlü sıfat-fiil kalıplarını arar. "
        f"Ham endeks (klasik ABG tanımı) 1 + (şahin−güvercin)/(şahin+güvercin) biçimindedir ve "
        f"az sayıda eşleşmede uçlara yapışma eğilimindedir (ör. 1 şahin/0 güvercin de, "
        f"20 şahin/0 güvercin de 2.00 verir). Bu raporda ASIL gösterge, paydaya bir düzeltme "
        f"sabiti eklenen (K={utils.ABG_SHRINK_K:.0f}) YUMUŞATILMIŞ endekstir: eşleşme sayısı "
        f"azaldıkça skor nötre (1.0) çekilir, arttıkça ham orana yaklaşır. Her iki değer de "
        f"özet tabloda birlikte gösterilir; ayrıca kaç eşleşmeye (n_match) dayandığı da raporlanır "
        f"— düşük n_match, endeksin güvenilir biçimde yorumlanamayacağının işaretidir."
    )
    _add_heading(doc, "2.2 CB-RoBERTa — cümle bazlı model, ton ve rejim ayrımı", level=2)
    doc.add_paragraph(
        "Şahin/güvercin sınıflandırması, TCMB PPK metinleri üzerinde eğitilmiş bir RoBERTa "
        "modeli (mrince/CBRT-RoBERTa-HawkishDovish-Classifier) ile CÜMLE düzeyinde yapılır "
        "(tam metin tek seferde verilirse hem 512 token sınırına takılır hem de literatürdeki "
        "yöntemle [Apel & Blix Grimaldi; Picault & Renault] tutarsız olurdu). Her cümle için "
        "ton = P(Şahin) − P(Güvercin) hesaplanır; kararı doğrudan bildiren cümle "
        f"({utils.DOC_ACTION_WEIGHT:.0f}× ağırlıkla) diğerlerinden daha belirleyicidir. "
        f"Cümle bazında şahin/güvercin/nötr ayrımı ±{utils.DOC_STANCE_DEADBAND:.2f}'lik tek bir "
        "eşikle (deadband) yapılır ve bu rapordaki TÜM cümle sayımları bu eşiği kullanır."
    )
    doc.add_paragraph(
        "TON (o dönemin ham/kalibre skoru) ile REJİM (🦅/🕊️/⚖️ etiketi) FARKLI şeylerdir: "
        "ton her dönem yeniden hesaplanan bir sayıdır; rejim ise robust z-skor → tanh kalibrasyonu "
        "→ üstel hareketli ortalama (EMA) → HİSTEREZİS BANDI ile yumuşatılmış bir etikettir. "
        "Histerezis, ardışık dönemler arasında gürültüden kaynaklı ani rejim sıçramalarını "
        "engeller: skor güçlü bir eşiği aşana kadar önceki rejim korunur, yalnızca ters yöndeki "
        "orta bir eşiği de geçerse rejim değişir. Bu yüzden bir dönemin HAM tonu şahine yakın "
        "olsa bile, rejim etiketi hâlâ 'Nötr' görünebilir — bu bir hata değil, tasarımdır."
    )
    _add_heading(doc, "2.3 Konu (tema) sınıflandırması", level=2)
    doc.add_paragraph(
        "Konu ataması MODEL TABANLI DEĞİL, düzenli-ifade (regex) sözlüğü tabanlıdır: her cümle "
        f"{len(utils.THEME_ORDER)-1} tanımlı temadan ({', '.join(utils.THEME_ORDER[:-1])}) "
        "hangisine ait anahtar kalıpları en çok içeriyorsa o temaya atanır (kararı bildiren "
        "cümle her zaman doğrudan 'Politika Duruşu' sayılır). Tek-etiket görünüm metnin "
        "KOMPOZİSYONUNU (%100'e tamamlanan pay), çok-etiketli görünüm ise KAPSAMINI (bir cümle "
        "birden çok temaya değinebilir, toplam %100'ü aşabilir) yansıtır; bu raporda ısı "
        "haritaları kompozisyon (tek etiket) moduyla üretilmiştir."
    )
    _add_heading(doc, "2.4 Metin-içi konum analizi", level=2)
    doc.add_paragraph(
        "Her karar metni kendi cümle sayısına göre eşit bölümlere ayrılır (bu raporda varsayılan "
        "3 bölüm: Giriş / Gövde / Kapanış) ve her bölümün ortalama tonu hesaplanır. Amaç, iletişimin "
        "'anlatı mimarisini' görmektir — ör. temkinli bir açılışın ardından sıkı bir çerçeve "
        "bloğu ve kararlı bir kapanış gelmesi, homojen bir tondan farklı bir strateji anlamına gelir."
    )
    _add_heading(doc, "2.5 Sınırlamalar", level=2)
    doc.add_paragraph(
        "(i) Sözlük ve model İngilizce metin üzerinde çalışır; analiz TCMB'nin kurumsal İngilizce "
        "çevirisi üzerinden yapılmıştır — çeviri tercihleri orijinal Türkçe metnin tonunu az da "
        "olsa kaydırabilir. (ii) Okunabilirlik skoru (Flesch) İngilizce dilbilgisi kurallarına "
        "göre tasarlanmıştır, Türkçe için doğrudan geçerli değildir. (iii) §7 Model Performansı "
        "bölümündeki tahmin modeli küçük bir örneklemle (aylık/6 haftalık PPK toplantıları) "
        "eğitilir; walk-forward doğrulama bile küçük örneklemde geniş belirsizlik taşır."
    )

    # =========================================================================
    # 3. ŞAHİN-GÜVERCİN DURUŞ ANALİZİ
    # =========================================================================
    _add_heading(doc, "3. Şahin-Güvercin Duruş Analizi", level=1)
    _add_takeaway(doc,
        f"Bu dönem ({donem}) ABG yöntemi metni \"{_abg_word.upper()}\", CB-RoBERTa modeli "
        f"\"{_ai_word.upper()}\" olarak okuyor"
        + (" — iki yöntem AYNI YÖNDE." if _abg_word == _ai_word and _abg_word != "belirsiz"
           else " — yöntemler FARKLI YÖNE işaret ediyor, aşağıdaki §3.3'e bakın.") +
        " Alt bölümler bu sonuca nasıl ulaşıldığını (§3.1–§3.2) ve önceki döneme göre "
        "kelime düzeyinde ne değiştiğini (§3.4) gösterir.")

    _add_heading(doc, "3.1 ABG (2019) sonucu", level=2)
    if text_now:
        s_abg, h_cnt, d_cnt, h_list, d_list, h_ctx, d_ctx, _ = utils.run_full_analysis(text_now)
        doc.add_paragraph(
            f"Bu dönemde ABG yöntemi {h_cnt} şahin, {d_cnt} güvercin modifikatör eşleşmesi "
            f"bulmuştur (yumuşatılmış endeks: {s_abg:.3f})."
        )
        ex_rows = []
        for term, sents in list(h_ctx.items())[:3]:
            ex_rows.append({"Yön": "🦅 Şahin", "Terim": term, "Örnek cümle": sents[0][:220]})
        for term, sents in list(d_ctx.items())[:3]:
            ex_rows.append({"Yön": "🕊️ Güvercin", "Terim": term, "Örnek cümle": sents[0][:220]})
        if ex_rows:
            _df_to_table(doc, pd.DataFrame(ex_rows))
    else:
        _add_note(doc, "Bu döneme ait ham metin bulunamadı.")

    _add_heading(doc, "3.2 CB-RoBERTa sonucu ve tarihsel bağlam", level=2)
    if ai_row is not None:
        doc.add_paragraph(
            f"Ton (Diff H-D): {ai_row.get('Diff (H-D)', float('nan')):+.3f}  ·  "
            f"Duruş: {ai_row.get('Duruş','—')}  ·  "
            f"Kalibre skor (EMA): {ai_row.get('AI Score (EMA)', float('nan')):.1f}  ·  "
            f"Rejim: {ai_row.get('AI Rejim', ai_row.get('Rejim','—'))}"
        )
    if ai_df is not None and not ai_df.empty:
        try:
            ai_sorted = ai_df.sort_values("period_date")
            fig_ai = (utils.create_tone_action_chart(ai_sorted, step=3)
                      if hasattr(utils, "create_tone_action_chart") and
                         {"Delta BP", "Aksiyon Yön"}.issubset(ai_sorted.columns)
                      else None)
            if fig_ai is not None:
                # utils.create_tone_action_chart alt paneldeki (Δ bp) y-eksenine PAY BIRAKMAZ —
                # en yüksek/en düşük bar, çizim alanının tam kenarına yapışır ve "kırpılmış"
                # görünür. utils.py'ye dokunmadan, burada döndürülen fig'e %18 boşluk payı ekliyoruz.
                _bp_vals = pd.to_numeric(ai_sorted.get("Delta BP"), errors="coerce").dropna()
                if not _bp_vals.empty:
                    _bp_max = float(_bp_vals.abs().max())
                    _pad = max(_bp_max * 0.18, 25.0)
                    _lo = min(0.0, float(_bp_vals.min()) - _pad)
                    _hi = max(0.0, float(_bp_vals.max()) + _pad)
                    fig_ai.update_yaxes(range=[_lo, _hi], row=2, col=1)
                _add_figure(doc, fig_ai, png_height=850,
                    caption="Üst panel: ton (EMA + histerezis); işaretçi ŞEKLİ gerçekleşen aksiyonu gösterir "
                            "(▲ artış, ▼ indirim, ■ sabit). Alt panel: gerçekleşen faiz değişimi (Δ bp). "
                            "Yukarıda duran bir ▼ 'şahin bir çerçeveyle iletilen indirim', aşağıda duran "
                            "bir ▲ 'güvercin bir çerçeveyle iletilen artış' anlamına gelir.")
            else:
                fig_ai = utils.create_ai_trend_chart_step(ai_sorted, step=3)
                _add_figure(doc, fig_ai, caption="CB-RoBERTa kalibre skoru (EMA) — tarihsel seyir, histerezis rejimiyle birlikte.")
        except Exception as e:
            _add_note(doc, f"Tarihsel grafik üretilemedi: {e}")
    else:
        _add_note(doc, "CB-RoBERTa önbelleğinde tarihsel seri bulunamadı.")

    _add_heading(doc, "3.3 Yöntemler Arası Karşılaştırma", level=2)
    doc.add_paragraph(_reconciliation_text(abg_row, sent_ozet))

    _add_heading(doc, "3.4 Sözcük ve İfade Değişimi (önceki döneme göre)", level=2)
    if text_now and text_prev:
        shift_rows = _lexical_shift_rows(text_now, text_prev, top_k=14)
        doc.add_paragraph(
            "Aşağıdaki tablo, önceki metne göre YENİ ORTAYA ÇIKAN, KAYBOLAN ya da belirgin biçimde "
            "GÜÇLENEN/ZAYIFLAYAN kelime ve iki-kelimelik ifadeleri listeler. Sayılar ham geçiş adedi "
            "değil, 1000 kelime başına orandır (‰) — böylece metin uzunluğu dönemden döneme "
            "değişse bile karşılaştırma anlamlı kalır. Bu, tek tek şahin/güvercin kelime sayımının "
            "(§2.1 ABG) ötesinde, iletişimin HANGİ SÖZCÜKLERİNİN öne çıktığını/geri çekildiğini gösterir."
        )
        if shift_rows:
            _df_to_table(doc, pd.DataFrame(shift_rows), font_size=8)
        else:
            _add_note(doc, "İki dönem arasında eşiği aşan belirgin bir sözcük/ifade değişimi tespit edilmedi.")
        _add_note(doc, "Bu liste anahtar-kelime frekansına dayanır; bir ifadenin şahin/güvercin "
                       "OLDUĞU anlamına gelmez — sadece önceki metne göre daha çok ya da daha az "
                       "kullanıldığı anlamına gelir. Yön yorumu için §2.1/§3.1 (ABG) ve §4 (cümle "
                       "bazlı ton) ile birlikte okunmalıdır.")
    else:
        _add_note(doc, "Karşılaştırma için önceki dönemin metni bulunamadığından bu bölüm atlandı.")

    # =========================================================================
    # 4. CÜMLE BAZLI TON DÖKÜMÜ
    # =========================================================================
    _add_heading(doc, "4. Cümle Bazlı Ton Dökümü", level=1)
    _add_takeaway(doc,
        "Dönemin özet skoru TEK bir sayıya sığmayabilir — metin genelde şahin VE güvercin "
        "cümlelerin bir karışımıdır. Bu bölüm metni cümle cümle açar: her cümlenin arka "
        "planı kendi tonuna göre (kırmızı=şahin, mavi=güvercin, gri=nötr) renklendirilir, "
        "böylece 'ortalama ton' rakamının ARKASINDA hangi ifadelerin durduğu görülür.")
    if not sent_period.empty:
        m1, m2, m3, m4 = sent_ozet.get("n", 0), sent_ozet.get("n_hawk", 0), sent_ozet.get("n_neut", 0), sent_ozet.get("n_dove", 0)
        doc.add_paragraph(f"Toplam {m1} cümle  ·  🦅 {m2} şahin  ·  ⚪ {m3} nötr  ·  🕊️ {m4} güvercin "
                          f"(ortalama ton: {sent_ozet.get('ort', float('nan')):+.3f})")
        try:
            fig_strip = utils.chart_sentence_strip(sent_period.sort_values("sent_idx"),
                                                    title=f"{donem} — cümle sırasına göre ton")
            _add_figure(doc, fig_strip, caption="Cümle sırasına göre ton (H−D). Metnin başından sonuna doğru okunur.")
        except Exception as e:
            _add_note(doc, f"Cümle-şeridi grafiği üretilemedi: {e}")

        _add_heading(doc, "Tüm metin — cümle cümle ton (renkli okuma)", level=2)
        doc.add_paragraph(
            "Aşağıda dönemin TAMAMI, her cümlenin arka planı kendi tonuna göre renklendirilmiş "
            "biçimde yer alır — tablodaki 6 örnek yerine metnin tamamını BAĞLAMIYLA birlikte "
            "okumak için."
        )
        _add_sentence_highlight_paragraph(doc, sent_period, utils.DOC_STANCE_DEADBAND)

        _add_heading(doc, "En şahin ifadeler", level=2)
        _df_to_table(doc, sahin_df.rename(columns={
            "sent_idx": "#", "sentence": "Cümle", "diff": "Ton",
            "agent_label": "İlgili Kesim", "theme_label": "Konu"}))
        _add_heading(doc, "En güvercin ifadeler", level=2)
        _df_to_table(doc, guvercin_df.rename(columns={
            "sent_idx": "#", "sentence": "Cümle", "diff": "Ton",
            "agent_label": "İlgili Kesim", "theme_label": "Konu"}))
    else:
        _add_note(doc, "Bu dönem için cümle önbelleği (roberta_sentences) bulunamadı. "
                       "Uygulamadaki «🗺️ Ton Haritası & Konular» sekmesinden önbelleği doldurduktan "
                       "sonra raporu yeniden üretin.")

    # =========================================================================
    # 5. KONU BAZLI ANALİZ
    # =========================================================================
    _add_heading(doc, "5. Konu Bazlı Analiz", level=1)
    _add_takeaway(doc,
        "Rapor buraya kadar metnin GENEL tonuna baktı; bu bölüm aynı soruyu KONU BAZINDA "
        "sorar — ör. metin enflasyon konusunda mı şahin, büyüme konusunda mı güvercin? "
        "İki farklı görsel var: ısı haritası (§5.1) 'hangi konuda ne kadar şahin/güvercin' "
        "sorusunu, konu kapsamı grafiği (§5.2) ise 'metin ZAMANLA hangi konulara daha çok "
        "yer ayırıyor' sorusunu cevaplar — ikisi FARKLI eksenler ölçer, karıştırılmamalı.")
    if not df_sent.empty:
        _add_heading(doc, "5.1 Konu × Ton Isı Haritası", level=2)
        doc.add_paragraph(
            "Nasıl okunur: satırlar konuları (enflasyon, büyüme, istihdam vb.), sütunlar "
            "dönemleri (PPK toplantılarını) gösterir. Her hücre, O DÖNEM O KONUYA değinen "
            "cümlelerin ORTALAMA tonudur — koyu kırmızıya kaydıkça o konuda şahin, koyu "
            "maviye kaydıkça güvercin bir dil kullanılmış demektir; beyaza yakın hücreler "
            "nötrdür. Gri hücre bir 'sıfır' değil, VERİ YOKLUĞUdur — o dönem metin o konuya "
            "hiç ya da yeterince değinmemiştir (bkz. min_n eşiği). Amaç: 'genel ton şahin' "
            "gibi tek bir cümlenin ARKASINDA hangi konunun bu sonucu sürüklediğini görmek "
            "— ör. genel ton nötr görünse bile enflasyon satırı koyu kırmızı olabilir."
        )
        tmat, cmat = utils.tone_matrix(df_sent, "theme_label", min_n=2)
        if not tmat.empty:
            fig_topic = utils.chart_tone_heatmap(tmat, "Konu × Ton (dönem bazında ortalama)", ylab="Konu", counts=cmat)
            _add_figure(doc, fig_topic, caption="Kırmızı = şahin, mavi = güvercin, beyaz = nötr; gri = o dönem o konuya değinilmedi (veri yok, sıfır değil).")
            if donem in tmat.index:
                doc.add_paragraph(f"{donem} döneminde konu bazlı tonlar: " +
                                  ", ".join(f"{c} {tmat.loc[donem, c]:+.2f}" for c in tmat.columns if pd.notna(tmat.loc[donem, c])))
        else:
            _add_note(doc, "Isı haritası için yeterli veri yok.")

        _add_heading(doc, "5.2 Konu Kapsamı (zaman içinde, %)", level=2)
        doc.add_paragraph(
            "Nasıl okunur: bu, ısı haritasından TAMAMEN FARKLI bir soru sorar — konunun "
            "TONUNU değil, metindeki YER KAPLAMINI (kaç cümlenin bu konudan bahsettiğini) "
            "ölçer. Her dönemde (yatay eksende), renkli bantların TOPLAMI her zaman %100'dür "
            "— yani grafik metnin o dönemki 'gündem dağılımını' gösterir. Bir bandın "
            "kalınlaşması 'bu konuya daha çok cümle ayrıldı' demektir; bu konunun ŞAHİN mi "
            "GÜVERCİN mi olduğuyla ilgisizdir (onu §5.1'deki ısı haritasından okuyun). "
            "Pratik kullanım: bir bandın aniden büyümesi/küçülmesi, TCMB'nin iletişimde "
            "önceliği bir konudan diğerine kaydırdığının erken bir işareti olabilir."
        )
        try:
            share = utils.share_timeseries(df_sent, "theme_label")
            if share is not None and not share.empty:
                order = share.mean().sort_values(ascending=False).index.tolist()
                share = share[[c for c in order if c in share.columns]]
                fig_share = utils.chart_share_area(share, "Konu kapsamı (%) — dönemler arası kompozisyon",
                                                    styles=getattr(utils, "THEME_STYLE", None))
                if fig_share is not None:
                    _add_figure(doc, fig_share,
                        caption="Yığılmış alan grafiği: her dönemde payların toplamı %100'dür. DİKKAT — "
                                "bir konunun payı, kendisi hiç değişmeden de düşebilir (başka bir konu "
                                "büyüdüğü için); mutlak hacim değil, KOMPOZİSYON okunur.")
                else:
                    _add_note(doc, "Konu kapsamı grafiği üretilemedi.")
            else:
                _add_note(doc, "Konu kapsamı serisi için yeterli veri yok.")
        except Exception as e:
            _add_note(doc, f"Konu kapsamı grafiği üretilemedi: {e}")

        dvg = utils.divergence_table(df_sent, "theme_label")
        _add_heading(doc, "5.3 Konu bazlı özet (tüm dönemler)", level=2)
        _df_to_table(doc, dvg.rename(columns={"theme_label": "Konu"}))
    else:
        _add_note(doc, "Konu analizi için cümle önbelleği gerekli; bulunamadı.")

    # =========================================================================
    # 6. METİN İÇİ KONUM ANALİZİ
    # =========================================================================
    _add_heading(doc, "6. Metin İçi Konum Analizi", level=1)
    _add_takeaway(doc,
        "Bir metin başında temkinli başlayıp sonunda sıkı bir dille bitebilir (ya da tersi) "
        "— tek bir ortalama ton bu 'anlatı yapısını' kaybeder. Bu bölüm metni Giriş / Gövde "
        "/ Kapanış olarak üçe böler ve her bölümün ayrı tonunu gösterir; böylece kararın "
        "HANGİ KISMININ mesajı taşıdığı (ör. kapanış cümlesi asıl sinyal mi?) görülür.")
    if not df_sent.empty:
        pmat, pcnt = utils.position_tone_matrix(df_sent, bins=3, min_n=1)
        if not pmat.empty:
            fig_pos = utils.chart_tone_heatmap(pmat, "Metin İçi Konum × Ton (Giriş / Gövde / Kapanış)",
                                               ylab="Konum", counts=pcnt)
            _add_figure(doc, fig_pos, caption="Her karar metni kendi uzunluğuna göre 3 eşit bölüme ayrılmıştır.")
            pos_text = _position_narrative(pmat, donem)
            doc.add_paragraph(pos_text)
        else:
            pos_text = None
            _add_note(doc, "Konum analizi için yeterli cümle bulunamadı.")
    else:
        pos_text = None
        _add_note(doc, "Konum analizi için cümle önbelleği gerekli; bulunamadı.")

    # =========================================================================
    # 7. MODEL PERFORMANSI / BACKTEST
    # =========================================================================
    _add_heading(doc, "7. Model Performansı (Metin → Faiz Kararı Backtest)", level=1)
    _add_takeaway(doc,
        "Önceki bölümler metnin TONUNU (şahin/güvercin) ölçtü; bu bölüm farklı, daha somut "
        "bir soruya bakar: bu iletişim tarzı, faizin gerçekte KAÇ BAZ PUAN değişeceğini "
        "önceden kestirmede ne kadar işe yarıyor? Aşağıda önce SIRADAKİ toplantı için "
        "modelin ima ettiği rakam, sonra bu tahminin ne kadar güvenilir olduğunu gösteren "
        "geçmiş performans (backtest) detayları yer alır.")

    # ---- SIRADAKİ TOPLANTI TAHMİNİ — önce sonuç, detay altta ----
    _add_heading(doc, "7.1 Sıradaki Toplantı İçin Model Tahmini", level=2)
    next_pred = _next_meeting_prediction(model_pack, donem, text_now, df_logs)
    if next_pred:
        _yon_renk = {"ARTIRIM": "C0392B", "İNDİRİM": "1F4E9C", "SABİT": "7F7F7F"}[next_pred["yon"]]
        _label_donem = ("sıradaki toplantı (canlı)" if next_pred["is_live_forecast"]
                        else f"{next_pred['actual_next_donem']} toplantısı")
        pred_rows = [{
            "label": f"Model tahmini — {_label_donem}",
            "value": f"{next_pred['pred_bp']:+.0f} bp  ({next_pred['yon']})",
            "prev": None, "color": _yon_renk,
        }, {
            "label": "Ampirik belirsizlik aralığı (geçmiş hata dağılımının %10–%90'ı)",
            "value": f"{next_pred['interval'][0]:+.0f} … {next_pred['interval'][1]:+.0f} bp",
            "prev": None, "color": None,
        }]
        if next_pred["actual_next"] is not None:
            _match = "aynı yönde" if (
                (next_pred["actual_next"] > 0 and next_pred["yon"] == "ARTIRIM") or
                (next_pred["actual_next"] < 0 and next_pred["yon"] == "İNDİRİM") or
                (next_pred["actual_next"] == 0 and next_pred["yon"] == "SABİT")
            ) else "FARKLI yönde"
            pred_rows.append({
                "label": f"Gerçekleşen ({next_pred['actual_next_donem']}, karşılaştırma amaçlı)",
                "value": f"{next_pred['actual_next']:+.0f} bp  — tahminle {_match}",
                "prev": None, "color": None,
            })
        _add_clean_summary_card(doc, pred_rows)
        if next_pred["is_live_forecast"]:
            _add_note(doc,
                "Bu tahmin, elde bulunan EN GÜNCEL yayımlanmış PPK metnini (bu raporun kendi "
                "dönemi) 'sözlü yönlendirme' sinyali olarak kullanır, çünkü bir sonraki (henüz "
                "toplanmamış) toplantının metni doğası gereği MEVCUT DEĞİLDİR — hiçbir metin "
                "modeli bunu gerçek anlamda göremez. Bu nedenle bu bir 'kehanet' değil, 'mevcut "
                "iletişim tarzı ve makro seyir sürerse modelin ne beklediği' sorusunun cevabıdır. "
                "Belirsizlik aralığı geniştir ve tek başına bir işlem/pozisyon kararının dayanağı "
                "olarak kullanılmamalıdır — bkz. §7.2'deki tarihsel isabet oranı.",
                label="Yöntem sınırlaması:")
        else:
            _add_note(doc,
                "Bu rapor GEÇMİŞE dönük üretildiği için, modelin o dönem için ima ettiği tahmin "
                "ile GERÇEKTE yaşanan karar burada yan yana konulmuştur — modelin canlı bir "
                "tahmininden değil, geçmişe dönük bir doğrulamadan söz ediyoruz.",
                label="Not:")
    else:
        _add_note(doc, "Sıradaki toplantı tahmini üretilemedi (backtest modeli eğitilmedi, "
                       "geçmiş özellik seti (df_hist) rapora iletilmedi ya da bu dönem için "
                       "metin bulunamadı).")

    _add_heading(doc, "7.2 Geçmiş Performans (Backtest Detayları)", level=2)
    doc.add_paragraph(
        "§7.1'deki tahmin NASIL üretildi? Model, ELİNDEKİ (geçmiş) PPK metinlerindeki kelime "
        "ve karakter dizilerinin ne sıklıkla "
        "kadar sık geçiyor VE tüm metinler arasında ne kadar AYIRT EDİCİ olduğu) sayısal "
        "özniteliklere çevirir; bunlara gecikmeli (bir önceki toplantılardan) makro/faiz "
        "değişkenlerini ekler ve bu birleşik öznitelik kümesiyle bir sonraki delta_bp'yi tahmin "
        "eden düzenlileştirilmiş (Ridge) bir regresyon eğitir. Düzenlileştirme, TF-IDF'in ürettiği "
        "çok sayıda (metindeki her benzersiz kelime/karakter dizisi kadar) öznitelikle, az sayıdaki "
        "PPK kararına AŞIRI UYUMU (overfitting) önlemek için gereklidir. Doğrulama YÜRÜYEN-PENCERE "
        "(walk-forward) yöntemiyle yapılır: her tahmin, SADECE o tarihten ÖNCEKİ kararlarla eğitilmiş "
        "bir modelden gelir (gelecekteki hiçbir veri modele sızmaz) — bu, ton skorlarının yalnızca "
        "iç-tutarlı değil, dışarıdan da (gerçekleşen kararla) doğrulanabilir olup olmadığının testidir."
    )
    _add_note(doc,
        "Metin öznitelikleri + düzenlileştirilmiş (Ridge/Elastic Net) regresyon + gecikmeli "
        "değişkenlerle birleştirme + yürüyen/genişleyen-pencere doğrulama şeması, merkez bankası "
        "metninden politika faizi DEĞİŞİMİNİ tahmin eden güncel akademik çalışmalarla aynı ailededir "
        "— örn. Duarte & Laurini (2026), Forecasting dergisinde yayımlanan çalışmada metin "
        "temsili olarak TF-IDF yerine transformer tabanlı yoğun gömme (embedding) + PCA kullanıyor, "
        "ancak gecikmeli politika dinamikleriyle birleştirip düzenlileştirilmiş regresyonla "
        "genişleyen-pencere (expanding-window) doğrulamasıyla Δpolitika-faizini (baz puan) tahmin "
        "etme YAPISI birebir aynıdır. TF-IDF'in kendisi de bu literatürde (FOMC/merkez bankası metni "
        "üzerinden faiz kararı tahmini) yaygın kullanılan, embedding'lere göre daha basit/klasik ama "
        "yerleşik bir metin temsilidir. Kısacası: buradaki YÖNTEM İSKELETİ (metin özniteliği + "
        "düzenlileştirme + gecikmeli dinamik + sızıntısız doğrulama) literatürle tutarlıdır; TF-IDF "
        "seçimi ise bu iskeletin daha basit/yorumlanabilir bir versiyonudur — R²/MAE/hit-rate "
        "değerleri (aşağıda) bu basit temsilin BU örneklemde ne kadar işe yaradığının doğrudan kanıtıdır.",
        label="Yöntemin bilimsel dayanağı:")
    if model_pack and model_pack.get("metrics"):
        metrics = model_pack["metrics"]
        hit = _hit_rate(model_pack.get("pred_df"))
        doc.add_paragraph(_backtest_narrative(metrics, hit))
        met_rows = [
            {"Gösterge": "Gözlem sayısı (n)", "Değer": metrics.get("n", "—")},
            {"Gösterge": "MAE (baz puan)", "Değer": f"{metrics.get('mae', float('nan')):.1f}"},
            {"Gösterge": "RMSE (baz puan)", "Değer": f"{metrics.get('rmse', float('nan')):.1f}"},
            {"Gösterge": "R²", "Değer": f"{metrics.get('r2', float('nan')):.3f}"},
        ]
        if hit and hit.get("n", 0) > 0:
            met_rows.append({"Gösterge": "Yön isabet oranı (hike/cut kararlarında)",
                             "Değer": f"%{hit['hit_rate']*100:.0f}  (n={hit['n']})"})
        _df_to_table(doc, pd.DataFrame(met_rows))

        pred_df = model_pack.get("pred_df")
        if pred_df is not None and not pred_df.empty:
            import plotly.graph_objects as go
            fig_bt = go.Figure()
            fig_bt.add_trace(go.Bar(x=pred_df["period_date"], y=pred_df["delta_bp"], name="Gerçekleşen (bp)", marker_color="#2F5496"))
            fig_bt.add_trace(go.Scatter(x=pred_df["period_date"], y=pred_df["pred_delta_bp"], name="Model Tahmini (bp)",
                                        mode="lines+markers", line=dict(color="#C0392B", width=2)))
            fig_bt.update_layout(title="Gerçekleşen vs. Tahmin Edilen Faiz Değişimi (bp)", height=480,
                                 margin=dict(t=50, b=40, l=40, r=20))
            _add_figure(doc, fig_bt, caption="Walk-forward backtest: her nokta yalnızca kendinden önceki kararlarla eğitilmiş bir modelin tahminidir (veri sızıntısı yoktur).")

        coef_df = model_pack.get("coef_df")
        if coef_df is not None and not coef_df.empty:
            top_pos = coef_df.sort_values("coef", ascending=False).head(8)
            top_neg = coef_df.sort_values("coef", ascending=True).head(8)
            _add_heading(doc, "En etkili kelimeler (model katsayıları)", level=2)
            comb = pd.concat([
                top_pos.assign(Yön="Faiz artışına işaret"),
                top_neg.assign(Yön="Faiz indirimine işaret"),
            ])[["Yön", "feature", "coef"]].rename(columns={"feature": "Kelime/İfade", "coef": "Katsayı"})
            _df_to_table(doc, comb)
    else:
        _add_note(doc, "Backtest modeli bu rapor için eğitilmedi veya yeterli etiketli gözlem "
                       "(delta_bp dolu karar, en az ~10 adet) yok.")

    # =========================================================================
    # 8. ANALİST NOTLARI / DIŞ DEĞERLENDİRMELER (opsiyonel, elle girilir)
    # =========================================================================
    _add_heading(doc, "8. Analist Notları / Dış Değerlendirmeler", level=1)
    _add_note(doc,
        "Bu bölüm isteğe bağlıdır ve otomatik üretilmez. Buraya kendi yorumunuzu ya da dışarıda "
        "(ör. bir sohbet asistanından) aldığınız bir okuma önerisini yapıştırabilirsiniz. Böyle bir "
        "girdi kullanıyorsanız şunu unutmayın: büyük dil modeli çıktıları PROMPT'A DUYARLIDIR ve "
        "TEK SEFERLİK/TEKRAR-ÜRETİLEMEZ olabilir — burada aşağıya yapıştırılan metin, doğrulanmış "
        "bir ölçüm değil, ek bir görüştür ve o şekilde okunmalıdır. BİRDEN FAZLA dış görüş "
        "kullanıyorsanız (ör. birkaç farklı sohbet asistanından alınan yorumlar), önce KENDİ "
        "1-2 cümlelik sentezinizi yazın (\"X ve Y aynı yönde, Z farklı görüşte çünkü ...\"), "
        "sonra tam metinleri ayrı ayrı; üç ayrı görüşü sentezlemeden art arda yapıştırmak okuyucuyu "
        "\"hangisi doğru?\" sorusuyla baş başa bırakır. Ayrıca bu metinler BAĞIMSIZ DOĞRULAMA değil, "
        "NİTEL YORUM KONTROLÜ olarak adlandırılmalıdır — aradaki fark önemlidir.",
        label="Kullanım notu:")
    if analyst_note and analyst_note.strip():
        p = doc.add_paragraph()
        p.add_run(analyst_note.strip())
        _abg_word_for_note = _stance_word((abg_row["abg_index"] - 1.0) if abg_row is not None else None, 0.05)
        _ai_word_for_note = _stance_word(ai_row.get("Diff (H-D)"), utils.DOC_STANCE_DEADBAND) if ai_row is not None else "belirsiz"
        _caveat = _note_consistency_caveat(analyst_note, _abg_word_for_note, _ai_word_for_note)
        if _caveat:
            _add_note(doc, _caveat, label="⚠ Otomatik tutarlılık kontrolü:")
    else:
        doc.add_paragraph("(Bu rapor üretilirken bir analist notu girilmedi.)").italic = True

    # =========================================================================
    # 9. SONUÇ VE DEĞERLENDİRME
    # =========================================================================
    _add_heading(doc, "9. Sonuç ve Değerlendirme", level=1)
    backtest_metrics = model_pack.get("metrics") if model_pack else None
    hit = _hit_rate(model_pack.get("pred_df")) if model_pack else None
    pos_short = _position_short(pmat, donem) if 'pmat' in dir() and pmat is not None and not pmat.empty else None
    doc.add_paragraph(_conclusion(donem, abg_row, ai_row, sent_ozet, backtest_metrics, hit, pos_short))

    # =========================================================================
    # 10. KISALTMALAR VE SÖZLÜK
    # =========================================================================
    _add_heading(doc, "10. Kısaltmalar ve Sözlük", level=1)
    glossary = [
        ("ABG", "Apel & Blix-Grimaldi (2019) — sözlük/kural temelli şahin-güvercin ölçüm yöntemi."),
        ("CB-RoBERTa", "Bu raporun kullandığı model: mrince/CBRT-RoBERTa-HawkishDovish-Classifier "
                       "(Hugging Face). Taban model FacebookAI/roberta-base'dir; TCMB PPK özet "
                       "metinlerinden çıkarılmış ~7.200 gerçek cümleyle 3 sınıf (hawkish/dovish/"
                       "neutral) için fine-tune edilmiştir. Genel amaçlı bir duygu (pozitif/negatif) "
                       "sınıflandırıcısı DEĞİLDİR ve Fed konuşmalarıyla eğitilmiş, benzer isimli başka "
                       "akademik modellerle karıştırılmamalıdır — bu, TCMB metinlerine özel bir modeldir."),
        ("Ton / Rejim", "Ton = o dönemin ham/kalibre skoru; Rejim = EMA + histerezis bandıyla yumuşatılmış, ani sıçramalara karşı dirençli etiket."),
        ("EMA", "Üstel hareketli ortalama (Exponentially Weighted Moving Average) — yakın dönemlere daha çok ağırlık veren yumuşatma yöntemi."),
        ("Histerezis bandı", "Rejim etiketinin değişmesi için skorun belirli bir eşiği aşması gerektiği; küçük dalgalanmalarda önceki rejimin korunmasını sağlayan mekanizma."),
        ("Deadband (eşik)", f"Bir cümlenin/dokümanın şahin, güvercin ya da nötr sayılması için ton skorunun aşması gereken sınır (bu raporda cümle düzeyinde ±{utils.DOC_STANCE_DEADBAND:.2f})."),
        ("n_match", "ABG yönteminde tespit edilen toplam şahin+güvercin kelime eşleşmesi sayısı; düşükse endeks güvenilir yorumlanamaz."),
        ("AOFM", "Ağırlıklı Ortalama Fonlama Maliyeti — TCMB'nin fiilen uyguladığı ortalama fonlama faizi (ilan edilen politika faizinden farklı olabilir)."),
        ("PKA / İYA / HBA 12 Ay Enflasyon Beklentisi", "Piyasa Katılımcıları / İmalat Sanayi Anketi / Hanehalkı Anketi kaynaklı, gelecek 12 aya dair enflasyon beklentisi serileri."),
        ("TÜFE", "Tüketici Fiyat Endeksi (CPI)."),
        ("delta_bp", "İki ardışık PPK kararı arasındaki politika faizi değişimi, baz puan cinsinden (100 bp = 1 puan)."),
        ("MAE / RMSE / R²", "Tahmin hatasının ortalama mutlak değeri / kök ortalama kare hatası / modelin açıkladığı varyans oranı — üçü de standart regresyon performans ölçütleridir."),
        ("Yön isabet oranı (hit-rate)", "Tahmin edilen faiz hareketinin yönünün (artış/indirim) gerçekleşenle örtüşme oranı; mevcut kod tabanında bulunmadığı için bu rapor için eklenmiştir."),
    ]
    gdf = pd.DataFrame(glossary, columns=["Terim", "Açıklama"])
    _df_to_table(doc, gdf, col_widths_in=[1.6, 4.9])

    # =========================================================================
    # EK: VERİ KALİTESİ
    # =========================================================================
    _add_heading(doc, "Ek: Veri Kalitesi", level=1)
    try:
        diag = utils.diagnose(df_logs)
        if diag is not None and not diag.empty:
            bad = diag[diag["durum"] != "taze"]
            doc.add_paragraph(
                f"Toplam {len(diag)} karar kaydından {len(diag) - len(bad)} tanesi cümle "
                f"önbelleğinde güncel ('taze'); {len(bad)} tanesi eksik/bayat/eski sürüm. "
                f"Bu raporun cümle-bazlı bölümleri (§4, §5, §6) yalnızca güncel önbelleği "
                f"olan dönemler için tam doğrudur."
            )
            if not bad.empty:
                _df_to_table(doc, bad.rename(columns={"period_date": "Dönem", "durum": "Durum"})[
                    ["Dönem", "source", "Durum"]].assign(Dönem=lambda d: pd.to_datetime(d["Dönem"]).dt.strftime("%Y-%m")))
        else:
            _add_note(doc, "Veri kalitesi teşhisi üretilemedi (önbellek tablosuna erişilemedi).")
    except Exception as e:
        _add_note(doc, f"Veri kalitesi teşhisi sırasında hata: {e}")

    # =========================================================================
    # EK: KELİME BULUTU (frekans bölümüne kısa görsel katkı)
    # =========================================================================
    if text_now:
        try:
            wc_fig = utils.generate_wordcloud_img(text_now)
            if wc_fig is not None:
                _add_heading(doc, "Ek: Kelime Bulutu (bu dönem)", level=1)
                _add_figure(doc, wc_fig, caption=f"{donem} dönemi karar metni — en sık geçen kelimeler.", is_mpl=True)
        except Exception:
            pass

    # ---- kaydet ----
    doc.save(out_path)
    _cleanup_tmp_files()
    return out_path


# =============================================================================
# 4. STREAMLIT İÇİNDEN ÇAĞRILACAK SARMALAYICI (canlı veri)
# =============================================================================

def generate_full_report_for_period(donem: str, analyst_note: Optional[str] = None,
                                     out_path: Optional[str] = None,
                                     birim: Optional[str] = None) -> str:
    """
    app.py içindeki yeni "📄 Rapor" sekmesinden çağrılır. Supabase/EVDS'e utils.py
    üzerinden (yani zaten çalışan Streamlit oturumunun secrets'ı ile) erişir.
    """
    import datetime

    df_logs = utils.fetch_all_data()
    if df_logs is None or df_logs.empty:
        raise ValueError("market_logs tablosunda kayıt bulunamadı.")

    df_events = utils.fetch_events()

    abg_df = utils.calculate_abg_scores(df_logs)

    min_d = pd.to_datetime(df_logs["period_date"]).min().date()
    max_d = datetime.date.today()
    df_market, _err = utils.fetch_market_data_adapter(min_d, max_d)

    ai_df = utils.trend_series_from_cache()

    df_sent = utils.fetch_sentences()

    model_pack = None
    if utils.HAS_ML_DEPS:
        try:
            df_td = utils.textasdata_prepare_df_hybrid_cpi(df_logs, df_market)
            if not df_td.empty and df_td["delta_bp"].notna().sum() >= 10:
                model_pack = utils.train_textasdata_hybrid_cpi_ridge(df_td)
                if model_pack:
                    # §7.1 "Sıradaki Toplantı İçin Model Tahmini" bu ham öznitelik geçmişine
                    # ihtiyaç duyar (utils.predict_textasdata_hybrid_cpi'nin df_hist argümanı).
                    model_pack["df_hist"] = df_td
        except Exception as e:
            print(f"[report_builder] Backtest modeli eğitilemedi (rapor bu bölüm olmadan devam ediyor): {e}")
            model_pack = None

    if out_path is None:
        out_path = f"ppk_rapor_{donem}.docx"

    return build_report(df_logs, df_events, df_market, abg_df, ai_df, df_sent,
                        donem=donem, model_pack=model_pack, analyst_note=analyst_note,
                        out_path=out_path, birim=birim)
